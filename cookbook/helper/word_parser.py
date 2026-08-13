"""Read a Word (.docx) recipe written to the household's conventions into a plain structure.

Pure by design (REQ-006 Decision 2): no Django, no models, no request, no database. It takes
document bytes and returns a :class:`WordRecipe`, or raises with a reason. Mapping the result
onto Recipe/Step/Ingredient is ``cookbook/integration/word.py``'s job, which keeps the Word
conventions out of the model-writing code and lets the parser be run standalone over a whole
folder of documents.

The documents are styled rather than free-form, which is what makes this tractable:

    Title          -> the recipe name
    Heading "Zutaten ..."     -> opens an ingredient section
    bulleted paragraph        -> one raw ingredient line, left for IngredientParser to split
    Heading "Zubereitung ..." -> opens an instruction section
    body paragraph            -> one instruction step, in document order

Style names are matched in English ('Title', 'Heading 1', 'List Paragraph') even though the
documents carry German style *ids* ('Titel', 'berschrift1', 'Listenabsatz' - Word strips
non-ASCII from ids). python-docx resolves a built-in style to its canonical English name, so
the same code reads both. The German words looked for in the heading *text* are document
content, not interface strings.

A paragraph's bullet never decides what it is; the section it sits in does. Documents in this
collection end with bulleted "Tipps" paragraphs after the instruction heading, and those are
steps, not ingredients.
"""
import os
import re
from dataclasses import dataclass, field
from io import BytesIO

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

#: Word puts pictures inside a drawing; the blip carries the relationship id of the media part.
DRAWING_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
RELATIONSHIP_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

#: image types Tandoor's handle_image can actually open - this is what drops the .wmf files
SUPPORTED_IMAGE_TYPES = ('.jpeg', '.jpg', '.png', '.gif', '.webp')

#: heading prefixes that open a section, lower-cased
INGREDIENT_HEADING = 'zutaten'
INSTRUCTION_HEADING = 'zubereitung'

#: German number words a yield phrase may spell out ('Zutaten fuer sechs Brezeln')
NUMBER_WORDS = {
    'ein': 1,
    'eine': 1,
    'einen': 1,
    'einem': 1,
    'eins': 1,
    'zwei': 2,
    'drei': 3,
    'vier': 4,
    'fuenf': 5,
    'fünf': 5,
    'sechs': 6,
    'sieben': 7,
    'acht': 8,
    'neun': 9,
    'zehn': 10,
    'elf': 11,
    'zwoelf': 12,
    'zwölf': 12,
}

#: articles a component name may be introduced by ('Zutaten fuer den Teig' -> 'Teig')
ARTICLES = ('der', 'die', 'das', 'den', 'dem', 'des', 'ein', 'eine', 'einen', 'einem')

DEFAULT_SERVINGS = 1


class WordDocumentSkipped(Exception):
    """A document that is deliberately not turned into a recipe. Always carries a reason."""

    def __init__(self, reason):
        self.reason = reason
        super().__init__(reason)


class UnparseableDocument(WordDocumentSkipped):
    """Shaped differently from the convention - refused rather than guessed at (Decision 6)."""


class EmptyDocument(WordDocumentSkipped):
    """Yields no ingredients and no steps at all (Decision 7)."""


@dataclass
class WordComponent:
    """One part of a recipe: its own ingredients followed by its own instructions.

    An unnamed component is the recipe itself; a named one is a 'Zutaten fuer X' section.
    """
    name: str = ''
    ingredients: list = field(default_factory=list)
    instructions: list = field(default_factory=list)


@dataclass
class WordRecipe:
    name: str = ''
    servings: int = DEFAULT_SERVINGS
    servings_text: str = ''
    components: list = field(default_factory=list)
    image: bytes = None
    image_filetype: str = None
    #: headings that were neither ingredients nor instructions, kept so a corpus run can see them
    ignored_headings: list = field(default_factory=list)

    @property
    def ingredient_count(self):
        return sum(len(c.ingredients) for c in self.components)

    @property
    def step_count(self):
        return sum(len(c.instructions) for c in self.components)


def _iter_blocks(document):
    """Yield the body's paragraphs and tables in document order."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, document)
        elif child.tag.endswith('}tbl'):
            yield Table(child, document)


def _style_name(paragraph):
    try:
        return paragraph.style.name or ''
    except Exception:
        return ''


def _is_heading(paragraph):
    return _style_name(paragraph).startswith('Heading')


def _is_bulleted(paragraph):
    properties = paragraph._p.pPr
    return properties is not None and properties.numPr is not None


def _normalize(text):
    """Collapse the whitespace Word scatters through a paragraph, including non-breaking spaces."""
    return ' '.join((text or '').replace('\xa0', ' ').split())


def _heading_phrase(heading_text, prefix):
    """The part of a heading that follows its keyword, stripped of decoration.

    'Zutaten fuer den Belag' -> 'fuer den Belag';  'Zutaten (fuer zwei Bleche)' -> 'fuer zwei Bleche'
    """
    phrase = heading_text[len(prefix):].strip()
    phrase = phrase.strip(':-– ').strip()
    if phrase.startswith('(') and phrase.endswith(')'):
        phrase = phrase[1:-1].strip()
    return phrase


def _strip_leading_word(phrase, words):
    first, _, rest = phrase.partition(' ')
    if first.lower() in words:
        return rest.strip()
    return phrase


def _leading_count(phrase):
    """Split a yield phrase into its count and the text after it.

    Returns ``(None, phrase)`` when nothing resolves, which is what leaves the count at its
    default while keeping the phrase as servings text.
    """
    phrase = _strip_leading_word(phrase, ('für', 'fuer', 'for'))

    match = re.match(r'^(\d+)\s*(.*)$', phrase)
    if match:
        return int(match.group(1)), match.group(2).strip()

    first, _, rest = phrase.partition(' ')
    if first.lower() in NUMBER_WORDS:
        return NUMBER_WORDS[first.lower()], rest.strip()

    return None, phrase


def _component_name(phrase):
    """The name a component heading gives its section: 'fuer den Belag' -> 'Belag'."""
    phrase = _strip_leading_word(phrase, ('für', 'fuer', 'for'))
    phrase = _strip_leading_word(phrase, ARTICLES)
    return phrase.strip()


def _matching_component(components, name):
    """The already-open component a 'Zubereitung X' heading refers to, if any.

    Documents that list every ingredient section first and only then every instruction
    section would otherwise pour one component's method into another's.
    """
    if not name:
        return None
    wanted = name.casefold()
    for component in components:
        if component.name and component.name.casefold() == wanted:
            return component
    return None


def _ingredient_heading_count(document):
    """How many ingredient headings the document has, needed before the walk starts.

    The first one is the yield heading unless it names a component, and it only names a
    component when there are further ingredient sections for it to be one of.
    """
    count = 0
    for block in _iter_blocks(document):
        if isinstance(block, Paragraph) and _is_heading(block) and _normalize(block.text).lower().startswith(INGREDIENT_HEADING):
            count += 1
    return count


def _first_image(document):
    """The first embedded picture in document order, as ``(bytes, filetype)``.

    Only types handle_image can open are considered, so the collection's .wmf drawings are
    passed over rather than attached and then failed on.
    """
    for blip in document.element.body.iter(f'{{{DRAWING_NS}}}blip'):
        relationship_id = blip.get(f'{{{RELATIONSHIP_NS}}}embed')
        if not relationship_id:
            continue
        try:
            part = document.part.related_parts[relationship_id]
        except KeyError:
            continue
        filetype = os.path.splitext(str(part.partname))[1].lower()
        if filetype in SUPPORTED_IMAGE_TYPES:
            return part.blob, filetype
    return None, None


def parse_docx(data, fallback_name=''):
    """Turn .docx bytes into a :class:`WordRecipe`.

    :param data: bytes, or any file-like object positioned at the start of the document
    :param fallback_name: used as the recipe name when the document carries no name of its own
    :raises UnparseableDocument: the ingredients are in a table instead of a bulleted list
    :raises EmptyDocument: no ingredients and no steps at all
    """
    document = docx.Document(BytesIO(data) if isinstance(data, (bytes, bytearray)) else data)

    recipe = WordRecipe()
    components = []
    current = None  # the component new content belongs to
    target = None  # where instructions go - usually `current`, but see _matching_component
    in_ingredients = False
    in_instructions = False
    saw_table = False
    leading_text = []  # body paragraphs before the first heading, in case there is no Title

    ingredient_headings = _ingredient_heading_count(document)
    seen_ingredient_headings = 0

    for block in _iter_blocks(document):
        if isinstance(block, Table):
            # never read for content (Decision 6) - only remembered, so the rejection below
            # can say *why* a document without a bulleted ingredient list was refused
            saw_table = True
            continue

        text = _normalize(block.text)
        style = _style_name(block)

        if style == 'Title':
            if not recipe.name:
                recipe.name = text
            continue

        if _is_heading(block):
            lowered = text.lower()

            if lowered.startswith(INGREDIENT_HEADING):
                seen_ingredient_headings += 1
                phrase = _heading_phrase(text, INGREDIENT_HEADING)
                name = ''

                if seen_ingredient_headings == 1:
                    count, rest = _leading_count(phrase)
                    if count is not None:
                        recipe.servings = count
                        recipe.servings_text = rest
                    elif ingredient_headings > 1:
                        name = _component_name(phrase)
                    else:
                        recipe.servings_text = rest
                else:
                    name = _component_name(phrase)

                current = WordComponent(name=name)
                components.append(current)
                target = current
                in_ingredients, in_instructions = True, False
                continue

            if lowered.startswith(INSTRUCTION_HEADING):
                if current is None:
                    current = WordComponent()
                    components.append(current)
                named = _matching_component(components, _component_name(_heading_phrase(text, INSTRUCTION_HEADING)))
                target = named or current
                in_ingredients, in_instructions = False, True
                continue

            # a heading that is neither - it ends the ingredient list but its own text is not
            # recipe content; record it so a corpus run can see what was passed over
            if text:
                recipe.ignored_headings.append(text)
            if current is None:
                current = WordComponent()
                components.append(current)
            target = current
            in_ingredients, in_instructions = False, True
            continue

        if not text:
            continue

        if in_ingredients:
            if _is_bulleted(block):
                current.ingredients.append(text)
            continue

        if in_instructions:
            # a bulleted paragraph here is a step like any other (the "Tipps" lists)
            target.instructions.append(text)
            continue

        leading_text.append(text)

    recipe.components = components

    if not recipe.name:
        recipe.name = leading_text[0] if leading_text else _normalize(fallback_name)

    if recipe.ingredient_count == 0:
        if saw_table:
            raise UnparseableDocument('ingredients are in a table, not a bulleted list')
        if recipe.step_count == 0:
            raise EmptyDocument('no ingredients and no instructions')

    recipe.image, recipe.image_filetype = _first_image(document)
    return recipe
