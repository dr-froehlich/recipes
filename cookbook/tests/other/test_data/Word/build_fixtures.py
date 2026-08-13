"""Generate the committed .docx fixtures the REQ-006 Word importer tests read.

Run from the repository root when a fixture needs to change::

    python cookbook/tests/other/test_data/Word/build_fixtures.py

The generated documents are committed; this script is the reviewable source they came from.
None of them is a document from the household collection - each is written here, to the
collection's conventions, so the gate never asserts against production content.

The one thing this does beyond calling python-docx is renaming the style ids to the German
ones Word actually writes ('Titel', 'berschrift1' - Word strips the non-ASCII, 'Listenabsatz').
The style *names* stay English, which is exactly the shape of the real corpus and the bridge
cookbook/helper/word_parser.py relies on: python-docx resolves a built-in style back to its
canonical English name whatever the id says. Fixtures written with python-docx's own English
ids would leave that bridge untested.
"""
import os
from io import BytesIO

import docx
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))

#: python-docx style id -> the id Word writes for the same built-in style in a German install
GERMAN_STYLE_IDS = {
    'Title': 'Titel',
    'TitleChar': 'TitelZchn',
    'Heading1': 'berschrift1',
    'Heading1Char': 'berschrift1Zchn',
    'ListParagraph': 'Listenabsatz',
}

#: attributes anywhere in the package that refer to a style by id
STYLE_REFERENCES = ('basedOn', 'next', 'link', 'pStyle', 'rStyle')


def bulleted(document, text):
    """A List Paragraph carrying paragraph-level numbering, the way Word writes a bullet."""
    paragraph = document.add_paragraph(text, style='List Paragraph')
    properties = paragraph._p.get_or_add_pPr()
    numbering = properties.makeelement(qn('w:numPr'), {})
    level = numbering.makeelement(qn('w:ilvl'), {qn('w:val'): '0'})
    number = numbering.makeelement(qn('w:numId'), {qn('w:val'): '1'})
    numbering.append(level)
    numbering.append(number)
    properties.append(numbering)
    return paragraph


def photograph(color=(200, 90, 40)):
    """A tiny JPEG standing in for a recipe photograph."""
    stream = BytesIO()
    Image.new('RGB', (48, 32), color).save(stream, format='JPEG')
    stream.seek(0)
    return stream


def germanize(path):
    """Rewrite the saved package so built-in styles carry the ids a German Word writes."""
    package = docx.Document(path)

    for style in package.styles.element.findall(qn('w:style')):
        current = style.get(qn('w:styleId'))
        if current in GERMAN_STYLE_IDS:
            style.set(qn('w:styleId'), GERMAN_STYLE_IDS[current])

    for part in (package.styles.element, package.element.body):
        for reference in STYLE_REFERENCES:
            for element in part.iter(qn(f'w:{reference}')):
                value = element.get(qn('w:val'))
                if value in GERMAN_STYLE_IDS:
                    element.set(qn('w:val'), GERMAN_STYLE_IDS[value])

    package.save(path)


def save(document, name):
    path = os.path.join(HERE, name)
    document.save(path)
    germanize(path)
    print(f'wrote {name}')


def canonical():
    """Every canonical signal in one document: title, yield, ingredients, steps, a component, a photo."""
    document = docx.Document()
    document.add_paragraph('Nussecken vom Blech', style='Title')

    heading = document.add_paragraph('Zutaten für 4 Personen', style='Heading 1')
    heading.add_run().add_picture(photograph(), width=Inches(1.5))
    for line in ('375 g Weizenmehl (550er)', '1 Tl Backpulver', '130 g Butter', '2 Eier'):
        bulleted(document, line)

    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Aus den Teigzutaten einen Knetteig verarbeiten.')
    document.add_paragraph('Den Backofen auf 175°C Ober-/Unterhitze vorheizen.')
    document.add_paragraph('Den Teig auf einem gefetteten Backblech ausrollen.')

    document.add_paragraph('Zutaten für die Füllung', style='Heading 1')
    for line in ('200 g gemahlene Haselnüsse', '3 El Wasser'):
        bulleted(document, line)

    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Die Nüsse mit dem Wasser verrühren.')
    document.add_paragraph('Die Masse auf den Teig geben und 25 Minuten backen.')

    save(document, 'canonical.docx')


def servings_number_word():
    """A yield spelled out in words, with a parenthetical the servings text has to keep."""
    document = docx.Document()
    document.add_paragraph('Laugenbrezeln', style='Title')
    document.add_paragraph('Zutaten für sechs Brezeln (ein Blech)', style='Heading 1')
    for line in ('170 g Weizenmehl (550er)', '14 g frische Hefe', '½ Tl Zucker'):
        bulleted(document, line)
    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Aus den Zutaten einen Hefeteig kneten.')
    save(document, 'servings_number_word.docx')


def servings_unresolvable():
    """A yield phrase with no number in it at all - the count stays at its default."""
    document = docx.Document()
    document.add_paragraph('Sonntagsbraten', style='Title')
    document.add_paragraph('Zutaten für die ganze Familie', style='Heading 1')
    for line in ('1 kg Rindfleisch', 'Salz'):
        bulleted(document, line)
    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Das Fleisch von allen Seiten scharf anbraten.')
    save(document, 'servings_unresolvable.docx')


def plain():
    """The ordinary shape: one ingredient section, one instruction section, no photograph."""
    document = docx.Document()
    document.add_paragraph('Kartoffelsalat', style='Title')
    document.add_paragraph('Zutaten', style='Heading 1')
    for line in ('500 g Kartoffeln', '1 Zwiebel', '3 El Essig'):
        bulleted(document, line)
    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Die Kartoffeln kochen, pellen und in Scheiben schneiden.')
    document.add_paragraph('Mit der Zwiebel und dem Essig vermengen und ziehen lassen.')
    save(document, 'plain.docx')


def table_ingredients():
    """The inverted shape pasted from a recipe site: ingredients in a table, steps as body text.

    Parsed with the canonical profile this would yield a recipe whose ingredient list is its
    method, which is the failure Decision 6 refuses to risk.
    """
    document = docx.Document()
    document.add_paragraph('Laugenbrezn')
    document.add_paragraph('Zutaten fürPortionen')

    table = document.add_table(rows=0, cols=2)
    for amount, food in (('170 g', 'Weizenmehl (550)'), ('⅔ TL', 'Salz'), ('14 g', 'frische Hefe')):
        cells = table.add_row().cells
        cells[0].text = amount
        cells[1].text = food

    document.add_paragraph('Aus den Zutaten einen Hefeteig kneten.')
    document.add_paragraph('Den Teig gehen lassen und zu Brezeln formen.')
    save(document, 'table_ingredients.docx')


def empty_headings():
    """Named and sectioned, but nothing under either heading."""
    document = docx.Document()
    document.add_paragraph('Angefangenes Rezept', style='Title')
    document.add_paragraph('Zutaten', style='Heading 1')
    document.add_paragraph('Zubereitung', style='Heading 1')
    save(document, 'empty_headings.docx')


def no_recipe():
    """Not a recipe at all - the collection's cover pages and blank templates look like this."""
    document = docx.Document()
    document.add_paragraph('Rezeptecover', style='Title')
    document.add_paragraph('Ein Deckblatt für den Ordner.')
    document.add_paragraph('')
    save(document, 'no_recipe.docx')


if __name__ == '__main__':
    canonical()
    servings_number_word()
    servings_unresolvable()
    plain()
    table_ingredients()
    empty_headings()
    no_recipe()
