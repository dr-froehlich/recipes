"""Generate the committed .docx fixtures the REQ-007 repair tests read.

Run from the repository root when a fixture needs to change::

    python cookbook/tests/other/test_data/Word/build_repair_fixtures.py

Same contract as ``build_fixtures.py`` beside it, whose helpers this reuses: the documents are
committed, this script is the reviewable source they came from, and none of them is a document
from the household collection. Every line below is written here, to the collection's
conventions, so the gate never asserts against production content.

Between them the two documents exercise every section of ``cookbook/integration/word_repairs.py``
- cut words, drop words, known units and their aliases, the unit conversion, keep-whole, food
aliases, both pre-parse fixes, grades, and the two guards that keep a repair from making a line
worse.
"""
import os
import sys

import docx

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from build_fixtures import bulleted, photograph, save  # noqa: E402  (needs HERE on the path)


def hefezopf():
    """Grades, the spelled-out unit, the conversion, a displaced noun, a range, drop words."""
    document = docx.Document()
    document.add_paragraph('Hefezopf', style='Title')

    heading = document.add_paragraph('Zutaten für 4 Personen', style='Heading 1')
    heading.add_run().add_picture(photograph(), width=docx.shared.Inches(1.5))
    for line in (
        '500 g 1050er Weizenmehl',
        '1 Tasse Milch',
        '1 gehäufter Teelöffel Salz',
        '2 EL flüssiger Honig',
        'Etwas Zucker',
        '4–6 Eier',
        '1 Eigelb zum Bestreichen',
        '200 g Weizenmehl (550er)',
    ):
        bulleted(document, line)

    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Alle Zutaten zu einem glatten Teig verkneten.')
    document.add_paragraph('Den Teig zu einem Zopf flechten und 35 Minuten backen.')

    save(document, 'repairs_hefezopf.docx')


def wurstsalat():
    """Adjectives in the unit slot, both spellings of one food and one unit, keep-whole, Je."""
    document = docx.Document()
    document.add_paragraph('Wurstsalat mit Avocado', style='Title')

    document.add_paragraph('Zutaten für 2 Personen', style='Heading 1')
    for line in (
        '6 Wiener Würstchen',
        '2 reife Avocados',
        '1 Avocado',
        '3 Esslöffel Olivenöl',
        'Je 1 rote und gelbe Paprika',
        '1 Dose gehackte Tomaten',
        '2 Tomaten',
        'Saft einer Limette',
        '200 g geriebener Käse',
        'Öl für die Form',
    ):
        bulleted(document, line)

    document.add_paragraph('Zubereitung', style='Heading 1')
    document.add_paragraph('Alles in mundgerechte Stücke schneiden und vermengen.')

    save(document, 'repairs_wurstsalat.docx')


if __name__ == '__main__':
    hefezopf()
    wurstsalat()
